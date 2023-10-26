import tempfile
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor

from dynamic_document_update.convertor.document_key_mapper import document_key_mapper
from dynamic_document_update.text_data.texts import perimenopause, postmenopausal, bmi_under_18_5, \
    bmi_over_18_5_or_under_24_9, bmi_25_or_over, lose_weight_yes, lose_weight_no, energy_level, caffeine_1, \
    caffeine_2, caffeine_3, caffeine_4, willing_to_caffeine_free, willing_to_change_to_decaf, \
    alcohol_per_night_once_a_week, alcohol_everyday_per_night_in_week, alcohol_three_or_five_per_night_in_week, \
    alcohol_in_weekends, alcohol_one_or_three_per_night_in_week, water_1_5, water_6_8, water_9_12, \
    avoid_processed_foods_true, avoid_processed_foods_false, can_follow_weekly_meal_plans_yes, \
    can_follow_weekly_meal_plans_no, can_carry_out_food_preparation_yes, can_carry_out_food_preparation_no, \
    diet_type_meat_eater, diet_type_vegen, diet_type_vegetarian, diet_type_pescatarian, diet_type_fodmap, \
    willing_to_eat_gluten_free, willing_to_eat_wholefoods, willing_to_eat_dairy_free, \
    willing_to_eat_sugar_free, one_meal_per_day, two_meals_per_day, three_meals_per_day, \
    daily_meals_include_morning_tea, \
    daily_meals_include_afternoon_tea, food_intolerances_allergies, food_crave_dark_chocolate, \
    food_crave_sweet, food_crave_salt, food_crave_cheese_and_dairy, sugar_drink, challenging, restorative, \
    restorative_yoga_and_stretching, restorative_resistance_training, restorative_cardiovascular_exercise, \
    restorative_mindfulness, restorative_meditation, restorative_qigong, challenging_meditation, \
    challenging_mindfulness, challenging_qigong, challenging_cardiovascular_exercise, challenging_resistance_training, \
    challenging_yoga_stretching, current_exercise_pilates, current_exercise_strength_training, current_exercise_run, \
    current_exercise_yoga, morning_exercise_schedule, afternoon_and_evening_exercise_schedule, no_time_for_exercise, \
    exercise_time_per_day_30_40_minutes, exercise_time_per_day_1_hours, exercise_time_per_day_2_or_more_hours, \
    training_location_gym, training_location_home, training_location_both, personal_trainer, \
    personal_trainer_frequency_once_in_week, personal_trainer_frequency_twice_in_week, increase_flexibility_yes, \
    increase_flexibility_no, increase_strength_yes, increase_strength_no, increase_fitness_yes, increase_fitness_no, \
    physical_injuries, mindfulness_stress_management_plan_true, bedtime_before_10, bedtime_between_10_to_12, \
    bedtime_after_12, fall_asleep_easily_no, difficult_to_get_back_to_sleep, average_sleep_hours_less_than_5, \
    average_sleep_hours_between_5_to_7, wake_up_feeling_rested_no, challenging_to_unwind_after_work, \
    unwind_methods_gym, unwind_methods_tv, unwind_methods_walk, unwind_methods_social_media, unwind_methods_alcohol, \
    unwind_methods_meditation, unwind_methods_work_until_bed_time, mental_health_rating_low, \
    enjoy_walking_for_exercise_yes, wellbeing_changes_mindfulness, wellbeing_changes_fitness, \
    wellbeing_changes_flexibility, wellbeing_changes_nutrition, wellbeing_changes_stress_management, smoking_quite_true, \
    on_medication, digestive_health_problems, \
    diagnosed_health_concerns, under_guidance_of_naturopath_dietician_nutritionist, alcohol_free_nights_per_week

VEGETARIAN = "Vegetarian - A person who does not eat any animal products, poultry, fish or seafood." \
             " Includes: Lacto-ovo (eats eggs & dairy), Lacto (eats dairy), Ovo (eats eggs), Pollo vegetarians " \
             "(eats poultry)."
VEGAN = "Vegan - A person who excludes all animal products including honey, and eggs as well as poultry, " \
        "fish, dairy and seafood."
PESCATARIAN = "Pescatarian - A person who does not eat meat or poultry but eats fish and seafood."
MEAT = "Meat Eater - This includes all red meat, eggs and poultry."

alcohol_consumption = {
    "Once a week": 1,
    "1-3 nights a week": 3,
    "3-5 nights a week": 5,
    "Every day": 7,
    "Weekends only": 2
}


def replace_placeholder(text, data):
    for key, value in data.items():
        text = text.replace(f"{key}", value)
    return text


def update_doc_file(data: dict) -> None:
    """function will read the FFL_Customised_Program_Template.docx and update the values and generate new documents"""
    mapper_data, conditional_mapper_data = document_key_mapper(data)
    doc = Document("FFL_Customised_Program_Template.docx")

    for paragraph in doc.paragraphs:
        paragraph.text = replace_placeholder(paragraph.text, mapper_data)

    # Handle conditional text based
    for i, paragraph in enumerate(doc.paragraphs):
        paragraph.paragraph_format.space_before = Inches(0.15)
        paragraph.paragraph_format.space_after = Inches(0.15)
        if "#10" in paragraph.text:
            if conditional_mapper_data["#10"]:
                paragraph.text = "Congratulations on your pregnancy ☺"
            else:
                paragraph.text = ""
        if "#11" in paragraph.text:
            if conditional_mapper_data["#11"]:
                paragraph.text = perimenopause
            else:
                paragraph.text = ""
        if "#12" in paragraph.text:
            if conditional_mapper_data["#12"]:
                paragraph.text = postmenopausal
            else:
                paragraph.text = ""
        if "#16" in paragraph.text:
            if float(conditional_mapper_data["#16"]) <= 18.5:
                paragraph.text = bmi_under_18_5
            elif 18.5 < float(conditional_mapper_data["#16"]) <= 24.9:
                paragraph.text = bmi_over_18_5_or_under_24_9
            else:
                paragraph.text = bmi_25_or_over
        if "#17" in paragraph.text:
            if conditional_mapper_data["#17"]:
                paragraph.text = lose_weight_yes
            else:
                paragraph.text = lose_weight_no
        if "#20" in paragraph.text:
            if int(conditional_mapper_data["#20"]) < 6:
                paragraph.text = energy_level
            else:
                paragraph.text = ""
        if "#25" in paragraph.text:
            if int(conditional_mapper_data["#25"]) == 1:
                paragraph.text = caffeine_1
            elif int(conditional_mapper_data["#25"]) == 2:
                paragraph.text = caffeine_2
            elif int(conditional_mapper_data["#25"]) == 3:
                paragraph.text = caffeine_3
            elif int(conditional_mapper_data["#25"]) == 4:
                paragraph.text = caffeine_4
            else:
                paragraph.text = ""
        if "#26" in paragraph.text:
            if conditional_mapper_data["#26"]:
                paragraph.text = willing_to_caffeine_free
            else:
                paragraph.text = ""
        if "#27" in paragraph.text:
            if conditional_mapper_data["#27"]:
                paragraph.text = willing_to_change_to_decaf
            else:
                paragraph.text = ""
        if "#28" in paragraph.text:
            if conditional_mapper_data["#28"] == "Once a week":
                paragraph.text = alcohol_per_night_once_a_week
            elif conditional_mapper_data["#28"] == "1-3 nights a week":
                paragraph.text = alcohol_one_or_three_per_night_in_week
            elif conditional_mapper_data["#28"] == "3-5 nights a week":
                paragraph.text = alcohol_three_or_five_per_night_in_week
            elif conditional_mapper_data["#28"] == "Every day":
                paragraph.text = alcohol_everyday_per_night_in_week
            elif conditional_mapper_data["#28"] == "Weekends only":
                paragraph.text = alcohol_in_weekends
            else:
                paragraph.text = ""
        if "#29" in paragraph.text:
            if alcohol_consumption[conditional_mapper_data["#29"]] <= alcohol_consumption[conditional_mapper_data["#28"]]:
                paragraph.text = alcohol_free_nights_per_week
            else:
                paragraph.text = ""
        if "#30" in paragraph.text:
            if conditional_mapper_data["#30"] == "1-5":
                paragraph.text = water_1_5
            elif conditional_mapper_data["#30"] == "6-8":
                paragraph.text = water_6_8
            elif conditional_mapper_data["#30"] == "9-12":
                paragraph.text = water_9_12
            else:
                paragraph.text = ""
        if "#31" in paragraph.text:
            if conditional_mapper_data["#31"]:
                paragraph.text = avoid_processed_foods_true
            else:
                paragraph.text = avoid_processed_foods_false
        if "#33" in paragraph.text:
            if conditional_mapper_data["#33"]:
                paragraph.text = can_follow_weekly_meal_plans_yes
            else:
                paragraph.text = can_follow_weekly_meal_plans_no
        if "#34" in paragraph.text:
            if conditional_mapper_data["#34"]:
                paragraph.text = can_carry_out_food_preparation_yes
            else:
                paragraph.text = can_carry_out_food_preparation_no
        if "#35" in paragraph.text:
            if conditional_mapper_data["#35"] == MEAT:
                paragraph.text = diet_type_meat_eater
            elif conditional_mapper_data["#35"] == VEGETARIAN:
                paragraph.text = diet_type_vegetarian
            elif conditional_mapper_data["#35"] == VEGAN:
                paragraph.text = diet_type_vegen
            elif conditional_mapper_data["#35"] == PESCATARIAN:
                paragraph.text = diet_type_pescatarian
            else:
                paragraph.text = diet_type_fodmap
        if "#36" in paragraph.text:
            food_willing_to_eat_mapper_data = conditional_mapper_data["#36"].split(",")
            food_willing_to_eat = ""
            for food in food_willing_to_eat_mapper_data:
                if food.strip() == "Wholefoods - Foods that come from nature.":
                    food_willing_to_eat += willing_to_eat_wholefoods
                elif food.strip() == "Gluten-free":
                    food_willing_to_eat += willing_to_eat_gluten_free
                elif food.strip() == "Dairy-free":
                    food_willing_to_eat += willing_to_eat_dairy_free
                elif food.strip() == "Sugar-free":
                    food_willing_to_eat += willing_to_eat_dairy_free
                elif food.strip() == "All of the above":
                    food_willing_to_eat += f"{willing_to_eat_wholefoods} \n " \
                                           f"{willing_to_eat_gluten_free} \n " \
                                           f"{willing_to_eat_dairy_free} \n" \
                                           f"{willing_to_eat_sugar_free}"
                else:
                    paragraph.text = ""
            paragraph.text = food_willing_to_eat
        if "#37" in paragraph.text:
            meals_per_mapper_data = conditional_mapper_data["#37"].split(",")
            meals_per_day = ""
            for meal in meals_per_mapper_data:
                if meal.strip() == "One":
                    meals_per_day += one_meal_per_day
                elif meal.strip() == "Two":
                    meals_per_day += two_meals_per_day
                elif meal.strip() == Three:
                    meals_per_day += three_meals_per_day
                else:
                    paragraph.text = ""
            paragraph.text = meals_per_day
        if "#38" in paragraph.text:
            daily_meal_mapper_data = conditional_mapper_data["#38"].split(",")
            daily_meal = ""
            for item in daily_meal_mapper_data:
                if item.strip() == "Morning Tea":
                    daily_meal += daily_meals_include_morning_tea
                elif item.strip() == "Afternoon Tea":
                    daily_meal += daily_meals_include_afternoon_tea
                # elif item.strip() == "Coffee":
                #     paragraph.text = willing_to_eat_dairy_free
                # elif item.strip() == "NIL":
                #     paragraph.text = willing_to_eat_dairy_free
                else:
                    paragraph.text = ""
            paragraph.text = daily_meal
        if "#39" in paragraph.text:
            if conditional_mapper_data["#39"]:
                paragraph.text = food_intolerances_allergies
            else:
                paragraph.text = ""
        if "#40" in paragraph.text:
            food_mapper_data = conditional_mapper_data["#40"].split(",")
            food_crave = ""
            for item in food_mapper_data:
                if item.strip() == "Salt":
                    food_crave += food_crave_salt
                elif item.strip() == "Sweet":
                    food_crave += food_crave_sweet
                elif item.strip() == "Dark chocolate":
                    food_crave += food_crave_dark_chocolate
                elif item.strip() == "Cheese and dairy":
                    food_crave += food_crave_cheese_and_dairy
                else:
                    paragraph.text = ""
            paragraph.text = food_crave
        if "#41" in paragraph.text:
            if conditional_mapper_data["#41"]:
                paragraph.text = sugar_drink
            else:
                paragraph.text = ""
        if "#43" in paragraph.text:
            if conditional_mapper_data["#43"] == "Restorative":
                paragraph.text = restorative
            else:
                paragraph.text = challenging
        if "#44" in paragraph.text:
            resistance_mapper_data = conditional_mapper_data["#44"].split(",")
            resistance_exercise = ""
            for exercise in resistance_mapper_data:
                if exercise.strip() == "Resistance training":
                    resistance_exercise += restorative_resistance_training
                elif exercise.strip() == "Yoga and stretching":
                    resistance_exercise += restorative_yoga_and_stretching
                elif exercise.strip() == "Qigong":
                    resistance_exercise += restorative_qigong
                elif exercise.strip() == "Cardiovascular exercise":
                    resistance_exercise += restorative_cardiovascular_exercise
                elif exercise.strip() == "Meditation":
                    resistance_exercise += restorative_meditation
                elif exercise.strip() == "Mindfulness":
                    resistance_exercise += restorative_mindfulness
                else:
                    paragraph.text = ""
            paragraph.text = resistance_exercise
        if "#45" in paragraph.text:
            challenging_mapper_data = conditional_mapper_data["#45"].split(",")
            challenging_exercise = ""
            for exercise in challenging_mapper_data:
                if exercise.strip() == "Resistance training":
                    challenging_exercise += challenging_resistance_training
                elif exercise.strip() == "Yoga and stretching":
                    challenging_exercise += challenging_yoga_stretching
                elif exercise.strip() == "Qigong":
                    challenging_exercise += challenging_qigong
                elif exercise.strip() == "Cardiovascular exercise":
                    challenging_exercise += challenging_cardiovascular_exercise
                elif exercise.strip() == "Meditation":
                    challenging_exercise += challenging_meditation
                elif exercise.strip() == "Mindfulness":
                    challenging_exercise += challenging_mindfulness
                else:
                    paragraph.text = ""
            paragraph.text = challenging_exercise
        if "#47" in paragraph.text:
            current_exercise_mapper_data = conditional_mapper_data["#47"].split(",")
            current_exercise = ""
            for exercise in current_exercise_mapper_data:
                if exercise.strip() == "I currently do yoga":
                    current_exercise += current_exercise_yoga
                elif exercise.strip() == "I currently do pilates":
                    current_exercise += current_exercise_pilates
                elif exercise.strip() == "I currently run":
                    current_exercise += current_exercise_run
                elif exercise.strip() == "I currently do strength training":
                    current_exercise += current_exercise_strength_training
                # @NOTE "not mentioned in sample docx"
                # elif exercise.strip() == "I walk daily":
                #     paragraph.text = challenging_meditation
                else:
                    paragraph.text = ""
            paragraph.text = current_exercise
        if "#48" in paragraph.text:
            exercise_time_mapper_data = conditional_mapper_data["#48"].split(",")
            exercise_time = ""
            for exercise in exercise_time_mapper_data:
                if exercise.strip() == "Morning":
                    exercise_time += morning_exercise_schedule
                elif exercise.strip() == "Afternoon" or exercise.strip() == "Evening":
                    exercise_time += afternoon_and_evening_exercise_schedule
                elif exercise.strip() == "I don't have time to exercise":
                    exercise_time += no_time_for_exercise
                else:
                    paragraph.text = ""
            paragraph.text = exercise_time
        if "#49" in paragraph.text:
            per_day_exercise_time_mapper_data = conditional_mapper_data["#49"].split(",")
            per_day_exercise_time = ""
            for exercise in per_day_exercise_time_mapper_data:
                if exercise.strip() == "30-45 minutes":
                    per_day_exercise_time += exercise_time_per_day_30_40_minutes
                elif exercise.strip() == "1 hour":
                    per_day_exercise_time += exercise_time_per_day_1_hours
                elif exercise.strip() == "2 or more hours":
                    per_day_exercise_time += exercise_time_per_day_2_or_more_hours
                else:
                    paragraph.text = ""
            paragraph.text = per_day_exercise_time
        if "#50" in paragraph.text:
            exercise_location_mapper_data = conditional_mapper_data["#50"].split(",")
            exercise_location = ""
            for location in exercise_location_mapper_data:
                if location.strip() == "Gym":
                    exercise_location += training_location_gym
                elif location.strip() == "Home":
                    exercise_location += training_location_home
                elif location.strip() == "Both":
                    exercise_location += training_location_both
                else:
                    paragraph.text = ""
            paragraph.text = exercise_location
        if "#51" in paragraph.text:
            if conditional_mapper_data["#51"]:
                paragraph.text = personal_trainer
            else:
                paragraph.text = ""
        if "#52" in paragraph.text:
            if conditional_mapper_data["#52"] == "Once a week":
                paragraph.text = personal_trainer_frequency_once_in_week
            elif conditional_mapper_data["#52"] == "Twice a week":
                paragraph.text = personal_trainer_frequency_twice_in_week
            else:
                paragraph.text = "\n \n"
        if "#55" in paragraph.text:
            if conditional_mapper_data["#55"]:
                paragraph.text = increase_flexibility_yes
            else:
                paragraph.text = increase_flexibility_no
        if "#57" in paragraph.text:
            if conditional_mapper_data["#57"]:
                paragraph.text = increase_strength_yes
            else:
                paragraph.text = increase_strength_no
        if "#59" in paragraph.text:
            if conditional_mapper_data["#59"]:
                paragraph.text = increase_fitness_yes
            else:
                paragraph.text = increase_fitness_no
        if "#60" in paragraph.text:
            if conditional_mapper_data["#60"]:
                paragraph.text = physical_injuries
            else:
                paragraph.text = ""
        if "#62" in paragraph.text:
            if conditional_mapper_data["#62"]:
                paragraph.text = mindfulness_stress_management_plan_true
            else:
                paragraph.text = ""
        if "#63" in paragraph.text:
            if conditional_mapper_data["#63"] == "Before 10pm":
                paragraph.text = bedtime_before_10
            elif conditional_mapper_data["#63"] == "Between 10-12pm":
                paragraph.text = bedtime_between_10_to_12
            elif conditional_mapper_data["#63"] == "After 12pm":
                paragraph.text = bedtime_after_12
            else:
                paragraph.text = ""
        if "#64" in paragraph.text:
            if not conditional_mapper_data["#64"]:
                paragraph.text = fall_asleep_easily_no
            else:
                paragraph.text = ""
        if "#66" in paragraph.text:
            if conditional_mapper_data["#66"]:
                paragraph.text = difficult_to_get_back_to_sleep
            else:
                paragraph.text = ""
        if "#67" in paragraph.text:
            if conditional_mapper_data["#67"] == "Less than 5 hours":
                paragraph.text = average_sleep_hours_less_than_5
            elif conditional_mapper_data["#67"] == "5-7 hours":
                paragraph.text = average_sleep_hours_between_5_to_7
            else:
                paragraph.text = ""
        if "#69" in paragraph.text:
            if not conditional_mapper_data["#69"]:
                paragraph.text = wake_up_feeling_rested_no
            else:
                paragraph.text = ""
        if "#71" in paragraph.text:
            if conditional_mapper_data["#71"]:
                paragraph.text = challenging_to_unwind_after_work
            else:
                paragraph.text = ""
        if "#72" in paragraph.text:
            unwind_methods_mapper_data = conditional_mapper_data["#72"].split(",")
            unwind_methods_text = ""
            for method in unwind_methods_mapper_data:
                if method.strip() == "Go to the gym":
                    unwind_methods_text += unwind_methods_gym
                elif method.strip() == "Go for a walk":
                    unwind_methods_text += unwind_methods_walk
                elif method.strip() == "Practice meditation":
                    unwind_methods_text += unwind_methods_meditation
                elif method.strip() == "Watch TV":
                    unwind_methods_text += unwind_methods_tv
                elif method.strip() == "Interact on social media":
                    unwind_methods_text += unwind_methods_social_media
                elif method.strip() == "Drink alcohol":
                    unwind_methods_text += unwind_methods_alcohol
                elif method.strip() == "I just keep working until I go to bed":
                    unwind_methods_text += unwind_methods_work_until_bed_time
                # elif method.strip() == "Hang out with my family/friends":
                #     paragraph.text = unwind_methods_work_until_bed_time
                else:
                    pass
            paragraph.text = unwind_methods_text
        if "#73" in paragraph.text:
            if conditional_mapper_data["#73"] <= 6:
                paragraph.text = mental_health_rating_low
            else:
                paragraph.text = ""
        if "#78" in paragraph.text:
            if conditional_mapper_data["#78"]:
                paragraph.text = enjoy_walking_for_exercise_yes
            else:
                paragraph.text = ""
        if "#80" in paragraph.text:
            wellbeing_changes_mapper_data = conditional_mapper_data["#80"].split(",")
            wellbeing_changes_text = ""
            for change in wellbeing_changes_mapper_data:
                if change.strip() == "Strength training":
                    wellbeing_changes_text += wellbeing_changes_fitness
                elif change.strip() == "Nutrition":
                    wellbeing_changes_text += wellbeing_changes_nutrition
                elif change.strip() == "Flexibility":
                    wellbeing_changes_text += wellbeing_changes_flexibility
                elif change.strip() == "Stress Management":
                    wellbeing_changes_text += wellbeing_changes_stress_management
                elif change.strip() == "Mindfulness":
                    wellbeing_changes_text += wellbeing_changes_mindfulness
                # elif change.strip() == "Meditation":
                #     wellbeing_changes_text += wellbeing_changes_mindfulness
                # elif change.strip() == "Increase daily exercise by walking":
                #     wellbeing_changes_text += unwind_methods_work_until_bed_time
                # elif change.strip() == "Remove processed foods":
                #     wellbeing_changes_text += unwind_methods_work_until_bed_time
                # elif change.strip() == "Remove my exposure to chemicals in products":
                #     wellbeing_changes_text += unwind_methods_work_until_bed_time
                else:
                    continue
            paragraph.text = unwind_methods_text
        if "#82" in paragraph.text:
            if conditional_mapper_data["#82"]:
                paragraph.text = smoking_quite_true
            else:
                paragraph.text = ""
        if "#85" in paragraph.text:
            if conditional_mapper_data["#85"]:
                paragraph.text = on_medication
            else:
                paragraph.text = ""
        if "#87" in paragraph.text:
            if conditional_mapper_data["#87"]:
                paragraph.text = digestive_health_problems
            else:
                paragraph.text = ""
        if "#89" in paragraph.text:
            if conditional_mapper_data["#89"]:
                paragraph.text = diagnosed_health_concerns
            else:
                paragraph.text = ""
        if "#91" in paragraph.text:
            if conditional_mapper_data["#91"]:
                paragraph.text = under_guidance_of_naturopath_dietician_nutritionist
            else:
                paragraph.text = ""

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        temp_file_path = temp_file.name
        doc.save(temp_file_path)

    return temp_file_path


def remove_extra_spaces(input_docx):
    doc = Document(input_docx)
    new_doc = Document()

    consecutive_blanks = 0

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            # Check for blank lines
            consecutive_blanks += 1
            if consecutive_blanks <= 1:
                new_doc.add_paragraph(paragraph.text)
        else:
            # For non-blank lines, reset the consecutive_blanks counter
            consecutive_blanks = 0
            new_doc.add_paragraph(paragraph.text)

    # Save the modified document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        temp_file_path = temp_file.name
        new_doc.save(temp_file_path)

    return temp_file_path


def add_image_to_docx(doc, image_path, width, height):
    first_paragraph = doc.add_paragraph()
    first_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = first_paragraph.add_run()
    run.add_picture(image_path, width=width, height=height)
    return first_paragraph


def convert_temp_file_to_proper_format(input_file, data):
    mapper_data, conditional_mapper_data = document_key_mapper(data)
    count = 0
    new_file = remove_extra_spaces(input_file)

    # Open the Word document
    doc = Document()

    # Add the image at the beginning
    add_image_to_docx(doc, 'images/fit_for_life.png', width=Inches(3.0), height=Inches(1.2))

    existing_doc = Document(new_file)

    # Append the content of the existing document
    for element in existing_doc.element.body:
        doc.element.body.append(element)

    # Define the name of the desired font (Helvetica)
    font_name = "Helvetica"

    for i, paragraph in enumerate(doc.paragraphs):
        paragraph.paragraph_format.space_after = Pt(0)
        if "#image1" in paragraph.text:
            if conditional_mapper_data["#35"] == VEGETARIAN:
                paragraph.clear()
                run = paragraph.add_run()
                run.add_picture('images/vegan_eater.png', width=Inches(3.5), height=Inches(2.3))
            else:
                paragraph.text = ""
        if "#image2" in paragraph.text:
            if conditional_mapper_data["#35"] == VEGAN:
                paragraph.clear()
                run = paragraph.add_run()
                run.add_picture('images/vegan_eater1.png', width=Inches(3.5), height=Inches(2.5))
            else:
                paragraph.text = ""
        if "#image3" in paragraph.text:
            if conditional_mapper_data["#35"] == VEGAN:
                paragraph.clear()
                run = paragraph.add_run()
                run.add_picture('images/vegan_eater2.png', width=Inches(3.7), height=Inches(2.5))
            else:
                paragraph.text = ""
        if "#fitForLife" in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture('images/fit_for_life.png', width=Inches(2.5), height=Inches(0.9))
        if paragraph.text == "":
            continue
        for run in paragraph.runs:
            if count < 3:
                if count == 0:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run.font.color.rgb = RGBColor(153, 0, 0)
                    run.font.size = Pt(20)
                else:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                count += 1

            run.font.name = font_name  # Set the font name
            if run.font.size is None:
                run.font.color.rgb = RGBColor(81, 81, 88)
                run.font.size = Pt(10.5)

    # Add the image at the end
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        temp_file_path = temp_file.name
        doc.save(temp_file_path)

    latest_created_doc = Document(temp_file_path)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        desired_file_name = f"{mapper_data.get('<Insert #1>')}.docx"
        new_file_path = os.path.join("/".join(temp_file.name.split("/")[0:-1]), desired_file_name)
        latest_created_doc.save(new_file_path)
    return new_file_path
